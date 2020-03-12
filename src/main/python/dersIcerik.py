from PyQt5.QtWidgets import QDialog, QTableWidget, QTableWidgetItem, QVBoxLayout, QHBoxLayout, QPushButton, QMessageBox, QLabel, QPlainTextEdit, QLineEdit
from PyQt5.QtCore import pyqtSlot, QTimer, QCoreApplication
from bs4 import BeautifulSoup
import os, sys
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
import re
import json
SUREMAX = 999999
ICERIKKLASOR = '\\icerik'

class dersIcerik(QDialog):
    def __init__(self, ctx):
        global debug
        super(dersIcerik, self).__init__()
        self.ctx = ctx
        debug = self.ctx.debug
        self.title = 'Ders İçerikleri'
        self.initUI()
        os.makedirs(self.ctx.anaKlasor + ICERIKKLASOR, exist_ok=True)
        if self.dersTabloAl():
            self.exec()

    def initUI(self):
        self.setWindowTitle(self.title)
        self.tableWidget = QTableWidget()
        self.ctx.dersSayisi=6
        self.tableWidget.setRowCount(self.ctx.dersSayisi)
        self.tableWidget.setColumnCount(5)
        self.tableWidget.setHorizontalHeaderLabels(['Ders Adı', 'Son Giriş Tarihi', 'Kalma Süresi', 'Gezinme Yüzdesi', 'Bağlantı'])
        self.tableWidget.setColumnWidth(0, 410)
        self.tableWidget.setColumnWidth(1, 110)
        self.tableWidget.setColumnWidth(2, 100)
        self.tableWidget.setColumnWidth(3, 100)
        self.tableWidget.setColumnWidth(4, 400)
        self.layout = QVBoxLayout()
        self.layout.addWidget(self.tableWidget)
        self.buttonsLayout= QHBoxLayout()
        self.layout.addLayout(self.buttonsLayout)
        self.btn_Update = QPushButton('İçerik Durumlarını Güncelle', self)
        self.btn_Update.clicked.connect(self.btnUpdateClicked)
        self.buttonsLayout.addWidget(self.btn_Update)
        self.btn_Kaydet = QPushButton('Seçilen Ders İçeriğini docx kaydet', self)
        self.btn_Kaydet.clicked.connect(self.btnKaydetClicked)
        self.buttonsLayout.addWidget(self.btn_Kaydet)
        self.btn_Baslat = QPushButton('Seçilen Ders İçin Otomatik İçerik Okumayı Başlat', self)
        self.btn_Baslat.clicked.connect(self.btnBaslatClicked)
        self.buttonsLayout.addWidget(self.btn_Baslat)
        self.setLayout(self.layout)
        self.setGeometry(50, 50, 1150, 300)

    def btnUpdateClicked(self):
        self.dersTabloAl(True)

    def btnBaslatClicked(self):
        if self.tableWidget.selectedItems()!=[]:
            dersler=self.dersler
            secilen= self.tableWidget.selectedItems()[0].row()
            if debug: self.ctx.logYaz(f"BaslatClicked: Icerik okunacak ders={dersler[secilen]['Ders']}")
            self.dersIcerikOku(secilen)

    def btnKaydetClicked(self):
        if self.tableWidget.selectedItems()!=[]:
            dersler=self.dersler
            secilen= self.tableWidget.selectedItems()[0].row()
            if debug: self.ctx.logYaz(f"BaslatClicked: Icerik kaydedilecek ders={dersler[secilen]['Ders']}")
            self.dersIcerikOku(secilen, Kaydet=True)

    def dersTabloAl(self, guncelle=False):
        durum = False
        if not guncelle:
            try:
                durum, dersler = self.dbDersDurumOku()
            except:
                durum = False
        if not durum:
            durum, dersler = self.dersDurumGetir()
        if durum:
            self.dersDurumDoldur(dersler)
            self.dersler = dersler
            return True
        else:
            return False

    def dersDurumDoldur(self, dersler):
        if len(dersler) != self.ctx.dersSayisi:
            self.ctx.dersSayisi = len(dersler)
            self.tableWidget.setRowCount(self.ctx.dersSayisi)
        i = 0
        for ders in dersler:
            self.tableWidget.setItem(i, 0, QTableWidgetItem(ders['Ders']))
            self.tableWidget.setItem(i, 1, QTableWidgetItem(ders['SonGiris']))
            self.tableWidget.setItem(i, 2, QTableWidgetItem(ders['KalmaSure']))
            self.tableWidget.setItem(i, 3, QTableWidgetItem(f"% {ders['Yuzde']:>5s}" ))
            self.tableWidget.setItem(i, 4, QTableWidgetItem(ders['Link']))
            i += 1
        # if ilkders > -1:
        #     if self.otomatik:
        #         self.tableWidget.item(ilkders, 2).setBackground(Qt.green)
        #     else:
        #         self.tableWidget.item(ilkders, 2).setBackground(Qt.white)
        if debug: self.ctx.logYaz(f'dersDurumDoldur: {i} ders dolduruldu')

    def dbDersDurumOku(self):
        dersler=[]
        if self.ctx.dbTableExists('derslerI'):
            cursor=self.ctx.dbCursor()
            cursor.execute('''SELECT ders, songiris, sure, yuzde, icerikno, link FROM derslerI''')
            i=0
            for row in cursor:
                dersler.append({'Ders': row[0]})
                dersler[i]['SonGiris'] = row[1]
                dersler[i]['KalmaSure'] = row[2]
                dersler[i]['Yuzde'] = row[3]
                dersler[i]['icerikno'] = row[4]
                dersler[i]['Link'] = row[5]
                i+=1
            if debug: self.ctx.logYaz(f'dersDurumOku: {i} ders veritabanından okundu')
            return True, dersler
        else:
            return False, dersler

    def dbYazDersler(self, dersler):
        #veritabanına yaz
        cursor=self.ctx.dbCursor()
        if self.ctx.dbTableExists('derslerI'):
            cursor.execute('''DROP TABLE derslerI''')
        cursor.execute('''CREATE TABLE IF NOT EXISTS
                        derslerI(ders TEXT, songiris TEXT, sure TEXT, yuzde TEXT, icerikno TEXT, link TEXT)''')
        self.ctx.db.commit()
        for ders in dersler:
            cursor.execute(f'''INSERT INTO derslerI(ders, songiris, sure, yuzde, icerikno, link)
                            VALUES('{ders['Ders']}','{ders['SonGiris']}','{ders['KalmaSure']}','{ders['Yuzde']}','{ders['icerikno']}','{ders['Link']}')''')
        self.ctx.db.commit()


    def dersDurumGetir(self):
        dersler = []
        durum = True
        mydosya=self.ctx.anaKlasor + '\\oys-ders.html'
        if not os.path.isfile(mydosya):
            if self.ctx.loginKontrol() is None:
                if debug: self.ctx.logYaz("dersDurumGetir: Giriş Yapılmadı, iptal?")
                self.ctx.login()
            cerezler = self.ctx.cerezOku()
            yanit = self.ctx.getSession().post(self.ctx.adres + '/ders_islemleri_ekran', cookies=cerezler)
            yanit.encoding = 'UTF-8'
            sayfa = yanit.text
            self.ctx.responseYaz(mydosya, sayfa)
        else:
            with open(mydosya, 'r', encoding="utf-8") as dosya:
                sayfa = dosya.read()
                dosya.close()
        try:
            soup = BeautifulSoup(sayfa, features='html.parser')
            div=soup.find('div',{'class':'col-md-12 tab-pane active', 'id':'contentHaftalikDers'})
            bulunandersler = div.find_all('div', {'class': 'card hover make-it-slow card-items'})
            i = 0
            for bulunanders in bulunandersler:
                eleman = bulunanders.find('button', {'class': 'btn btn-outline-linkedin lesson-state'})
                if eleman:
                    icerikno=eleman.attrs['onclick'].split('"')[1]
                    durum, ders = self.dersIcerikDrm(icerikno)
                else:
                    if debug: print(f"dersDurumGetir: eleman={eleman} bulunanders={bulunanders} bulunandersler=", bulunandersler)
                    break
                if not durum: break
                dersler.append({'Ders': ders[0].strip()})
                dersler[i]['SonGiris'] = ders[1].strip()
                dersler[i]['KalmaSure'] = ders[2].strip()
                dersler[i]['Yuzde'] = ders[3].strip()
                dersler[i]['icerikno'] = icerikno
                eleman = bulunanders.find('button', {'class': 'btn btn-outline-purple'})
                # print(eleman)
                icerikno=eleman.attrs['derskodu']
                durum, dersler[i]['Link'] = self.dersIcerikLink(icerikno)
                if debug: print(f"dersDurumGetir:{i} {dersler[i]}")
                i += 1
            self.dbYazDersler(dersler)
        except:
            if debug: print("dersDurumGetir: Hata var",  sys.exc_info()[0] )
            durum = False
        return durum, dersler

    def dersIcerikDrm(self, icerikno):
        ders = []
        veri = {'METHOD': 'GODRM', 'Mode': 'List', 'SID':icerikno, 'sg':''}
        self.ctx.onlineOl()
        cerezler=self.ctx.cerezOku()
        yanit = self.ctx.getSession().post(self.ctx.adres + '/ders_islemleri_ekran', data=veri, cookies=cerezler)
        yanit.encoding = 'UTF-8'
        sonuc = yanit.json()
        basarili = sonuc['Basarili']
        if basarili:
            # print('sonuc=', sonuc)
            soup = BeautifulSoup(sonuc['Deger'], features='html.parser')
            derstr = soup.find('tbody',{'id':'IcerikteKalmaSuresi'}).find('tr')
            if derstr:
                derstr = derstr.text
                # print(f"derstr={derstr}")
                # td = derstr.find('<td>')
                # while td>-1:
                #     derstr=derstr[td+4]
                #     print(f"derstr={derstr} td={td}")
                #     td = derstr.find('<td>')
                #     ders.append(derstr[td])
                ders=derstr.split('         ')
                durum = True
            else:
                ders[0]='derstr düzgün gelmedi...'
                durum = False
        else:
            ders[0]= 'yanit.json düzgün gelmedi...'
            durum = False
        if debug: print(f"dersIcerikDrm: {durum} <tr>'den split edilen ders={ders}")
        return durum, ders

    def dersIcerikLink(self, icerikno):
        ders = []
        veri = {'METHOD': 'GODIL', 'ID':icerikno, 'sg':''}
        self.ctx.onlineOl()
        cerezler=self.ctx.cerezOku()
        yanit = self.ctx.getSession().post(self.ctx.adres + '/ders_islemleri_ekran', data=veri, cookies=cerezler)
        yanit.encoding = 'UTF-8'
        sonuc = yanit.json()
        basarili = sonuc['Basarili']
        if basarili:
            # print('sonuc=', sonuc)
            soup = BeautifulSoup(sonuc['Deger'], features='html.parser')
            eleman = soup.find_all('a',{'class':'btn btn-xs btn-outline-purple'})
            if eleman:
                baglanti=eleman.attrs['onclick'].split("'")[1]
                # print(f"baglanti={baglanti}")
                durum = True
            else:
                baglanti='icerik düzgün gelmedi...'
                durum = False
        else:
            baglanti= 'yanit.json düzgün gelmedi...'
            durum = False
        if debug: print(f"dersIcerikLink: <a onClick=jQuery.openCourseContents({baglanti})")
        return durum, baglanti

    def dersIcerikOku(self, secilen, Kaydet=False):
        if debug: print(f"dersIcerikOku: secilen ders[{secilen}]={self.dersler[secilen]}")
        self.ctx.TimedMessageBox('dersIcerikOku',f"Seçtiğiniz ders ({ self.dersler[secilen]['Ders'] })\nLütfen Bekleyiniz...",QMessageBox.Ok, 3)
        self.ctx.IcerikDers = self.dersler[secilen]['Ders']
        mydosya=self.ctx.anaKlasor + f"{ICERIKKLASOR}\\oys-icerik-{self.dersler[secilen]['Ders'][:8]}.html"
        self.ctx.onlineOl()
        cerezler = self.ctx.cerezOku()
        if self.ctx.online or not os.path.isfile(mydosya):
            yanit = self.ctx.getSession().get(self.dersler[secilen]['Link'], cookies=cerezler)
            yanit.encoding = 'UTF-8'
            sayfa = yanit.text
            self.ctx.responseYaz(mydosya, sayfa)
        else:
            with open(mydosya, 'r', encoding="utf-8") as dosya:
                sayfa = dosya.read()
                dosya.close()
        soup = BeautifulSoup(sayfa, features='html.parser' )
        scripts = soup.find_all('script')
        for script in scripts:
            script2 = script.text
            script2 = script2.replace('lt;','<')
            script2 = script2.replace('gt;','>')
            v1 = 0
            v2 = script2.find('; var')
            while (v1 != -1):
                v1 = script2.find(';')
                if v1 != v2 :
                    script2 = script2[:v1] + ':' + script2[v1 + 1:]
                else:
                    v1 = -1
            if debug: print(f"dersIcerikOku: script={script2}")
            if 'var arrayData' in script2:
                script =  re.findall(r"=(.*?;)", script2 )
                if len(script)<2: script =  re.findall(r"=((?:\[.*\])?.*?;)", script2 )
                if debug: print(f"dersIcerikOku: len={len(script)} script[0]={script[0]}")
                jArray = script[0][:-1]
                if debug: print ("dersIcerikOku: jArray1=",jArray)
                try:
                    arrayData = json.loads(jArray)
                except:
                    self.ctx.TimedMessageBox('dersIcerikOku',"Hata: İçerik beklenildiği gibi değil",
                                             QMessageBox.Ok, 2)
                    arrayData = []
                if debug: print(f"dersIcerikOku: script[1]={script[1]}")
                jArray = script[1][1:-2]
                if debug: print("dersIcerikOku: jArray2=", jArray)
                if jArray=='null':
                    ogrStatus = []
                else:
                    try:
                        ogrStatus = json.loads(jArray)
                    except:
                        self.ctx.TimedMessageBox('dersIcerikOku', "Hata: İçerik beklenildiği gibi değil",
                                                 QMessageBox.Ok, 2)
                        ogrStatus = []
        if debug: print(f"dersIcerikOku: len={len(arrayData)} arrayData=", arrayData)
        if debug: print(f"dersIcerikOku: len={len(ogrStatus)} ogrStatus=", ogrStatus)
        veri = {'GMOD': 'Start'}
        yanit = self.ctx.getSession().post(self.dersler[secilen]['Link'], data=veri, cookies=cerezler)
        yanit.encoding = 'UTF-8'
        # if debug: print(yanit.text)
        sonuc = yanit.json()
        basarili = sonuc['Basarili']
        if basarili:
            # print("ok")
            ks, ss, dizin = self.dizinOlustur(arrayData)
            ss += ks
            if debug: print(f"dersIcerikOku: klasör={ks} sayfa={ss}")
            min = SUREMAX
            top = 0
            i = 0
            for o in ogrStatus:
                top += o['PDIH_ICERIK_SURE']
                for d in dizin:
                    if o['PDIH_ICERIK_MANIFEST_ADI']==d['ss']:
                        d['sure']=o['PDIH_ICERIK_SURE']
                        i = i + 1
                if o['PDIH_ICERIK_SURE'] < min:
                    min = o['PDIH_ICERIK_SURE']
            sayfalar=[]
            ii = 0
            for d in dizin:
                if d['sure'] <= (min+self.ctx.SureArtim) or (self.ctx.IcerikDS and d['ad'][:13]=='Değerlendirme') or self.ctx.IcerikTum or Kaydet:
                    sayfalar.append({'ad': d['ad'] })
                    sayfalar[ii]['link'] = f"{self.ctx.adres}/index.php?Reque=dersIcerikView&dersManifest={d['ss']}&manifestKey={d['kk']}&DP={d['dp']}&ss={ss}"
                    ii += 1
            if debug: print(f"dersIcerikOku: i={i} adet süre @ len(ogrStatus)={ len(ogrStatus) }. Top={timedelta(seconds = top)} min={min} okunacak sayfalar={ii}")
            if debug: self.ctx.logYaz(f"Okuma süresi {min+self.ctx.SureArtim} saniyeden az olan {ii} sayfa {self.ctx.SureArtim} sn.kadar okunacak...")
            if ii>0:
                self.ctx.cerezler = cerezler
                b = self.IcerikOkuma(self.ctx, sayfalar, Kaydet)
                durum, ders = self.dersIcerikDrm( self.dersler[secilen]['icerikno'] )
                if durum:
                    self.dersler[secilen]['SonGiris'] = ders[1].strip()
                    self.dersler[secilen]['KalmaSure'] = ders[2].strip()
                    self.dersler[secilen]['Yuzde'] = ders[3].strip()
                    self.dersDurumDoldur(self.dersler)
                    self.dbYazDersler(self.dersler)
                    if debug: print(f"dersIcerikOku: {self.dersler[secilen]['Ders']} güncellendi: SonGrş:{self.dersler[secilen]['SonGiris']} Yüzde={self.dersler[secilen]['Yuzde']}")
            else:
                self.ctx.TimedMessageBox('dersIcerikOku', f"{self.ctx.SureArtim}'dan az okunmuş ders bulunamadı", QMessageBox.Ok, 3)

    def dizinOlustur(self, arrayData):
        dizin = []
        ss = 0
        ks = 0
        i = -1
        for a in arrayData:
            i += 1
            elem = a[0]
            if debug and (i % 11)==0: print(f"dizinOlustur: i={i} elem={elem['text']} a.value={elem['value']}")
            dizin.append({ 'dp': elem['ders'] })
            dizin[i]['ss'] = elem['link']
            dizin[i]['sure'] = -1
            dizin[i]['ad'] = elem['text']
            dizin[i]['kk'] = elem['key']
            if elem['value'] is None:
                dizin[i]['tip'] = 'sayfa'
                ss += 1
            else:
                dizin[i]['tip'] = 'klasor'
                ks += 1
                k, s, alt = self.dizinOlustur(elem['value'])
                ks += k
                ss += s
                dizin += alt
                i += len(alt)
        # if debug: print(f"dizinOlustur: klasör={ks} sayfa={ss} i={i}")
        return ks, ss, dizin

    class IcerikOkuma(QDialog):
        def __init__(self, ctx, sayfalar, Kaydet):
            super(QDialog, self).__init__()
            self.ctx = ctx
            self.title = 'İçerik Okuma'
            self.initUI()
            self.sayfalar = sayfalar
            self.Kaydet = Kaydet
            self.sayfano = 0
            self.toplamsayfa = 0
            self.saniye = 0
            self.toplamsure = 0
            self.run = False
            if Kaydet: self.otomatikZamanla(1, self.IcerikKaydet)
            elif self.ctx.IcerikOto: self.otomatikZamanla(10, self.otomatikBasla) # 10 saniye sonra otomatik başla
            self.exec_()

        def initUI(self):
            self.setWindowTitle(self.title)
            self.lblDersAd = QLabel('Ders Adı', self)
            self.txtLink = QLineEdit('Ders Bağlantısı',self)
            self.btnBar = QHBoxLayout()
            self.lblStatus = QLabel('<Status>', self)
            self.btnGeri = QPushButton('Geri',self)
            self.btnBaslat = QPushButton('Başlat', self)
            self.btnIleri = QPushButton('İleri',self)
            self.btnBar.addWidget(self.lblStatus)
            self.btnBar.addWidget(self.btnGeri)
            self.btnBar.addWidget(self.btnBaslat)
            self.btnBar.addWidget(self.btnIleri)
            self.txtDers = QPlainTextEdit('İçerik', self)
            self.vlayout = QVBoxLayout()
            self.vlayout.addWidget(self.lblDersAd)
            self.vlayout.addWidget(self.txtLink)
            self.vlayout.addLayout(self.btnBar)
            self.btnGeri.clicked.connect(self.GeriClicked)
            self.btnBaslat.clicked.connect(self.BaslatClicked)
            self.btnIleri.clicked.connect(self.IleriClicked)
            self.vlayout.addWidget(self.txtDers)
            self.setLayout(self.vlayout)
            self.setGeometry(115, 115, 800, 511)

        @pyqtSlot()
        def closeEvent(self, event):
            if self.run:
                self.BaslatClicked()
            event.accept()

        def otomatikZamanla(self, sure, func):
            self.timerX = QTimer()
            self.timerX.timeout.connect(func)
            self.timerX.start(1000 * sure)

        def otomatikBasla(self):
            self.timerX.stop()
            self.timerX=None
            self.BaslatClicked()

        @pyqtSlot()
        def BaslatClicked(self):
            if self.timerX is not None:
                self.timerX.stop()
                self.timerX = None
            if not self.run:
                self.run = True
                self.timer = QTimer()
                self.timer.timeout.connect(self.IcerikOkuTimer)
                self.timer.start(1000)
                self.IcerikOkuTimer()
            else:
                self.run = False
                self.timer.stop()
                self.timer=None
                self.btnBaslat.setText(f"({(self.saniye % self.ctx.SureArtim) + 1})..{self.ctx.SureArtim} Devam Et")
                self.ctx.logYaz(f"BaslaClicked: {self.toplamsayfa} adet sayfa toplam {self.toplamsure} saniye  okundu.")
                self.otomatikZamanla(180, self.close) #oturum kapanabilir, 3 dakika sonra pencereyi kapat

        def sayacSifirla(self):
            self.saniye -= (self.saniye % self.ctx.SureArtim)

        def GeriClicked(self):
            if self.run:
                if self.sayfano > 1:
                    self.sayfano -= 2
                    self.sayacSifirla()

        def IleriClicked(self):
            if self.run:
                self.sayfano += 0
                self.sayacSifirla()

        def IcerikOkuTimer(self):
            if (self.saniye % self.ctx.SureArtim)==0:
                if self.sayfano < len(self.sayfalar):
                    self.IcerikOku(self.sayfano)
                self.sayfano += 1
                self.toplamsayfa += 1
                if self.sayfano > len(self.sayfalar):
                    self.sayfano = 0
                    self.BaslatClicked()
                    self.btnBaslat.setText('Tekrar Başlat')
                    if self.ctx.IcerikOto: self.close()
                    self.sayacSifirla()
            else:
                self.btnBaslat.setText(f"({(self.saniye % self.ctx.SureArtim) + 1})..{self.ctx.SureArtim} Durdur")
                self.setWindowTitle(f"{self.title} Geçen süre: {self.toplamsure+1}")
            self.saniye += 1
            self.toplamsure += 1

        def IcerikKaydet(self):
            klasor = f"{self.ctx.anaKlasor}{ICERIKKLASOR}"
            self.Belge = Document()
            self.Belge.add_heading(f"{self.ctx.IcerikDers} Ders İçeriği", 0)
            self.btnBar.setEnabled(False)
            self.timerX.stop()
            for no in range( len(self.sayfalar)):
                self.IcerikOku(no)
                if no % 10 == 0:  QCoreApplication.processEvents()
            dosya = f"{klasor}\\{self.ctx.IcerikDers[:8]}.docx"
            self.Belge.save(dosya)
            self.close()
            os.system(f'"{dosya}"')

        def IcerikOku(self, no):
            sayfalar = self.sayfalar
            yanit = self.ctx.getSession().get( sayfalar[no]['link'] , cookies=self.ctx.cerezler)
            yanit.encoding = 'utf-8'
            durum = yanit.status_code
            if durum==200:
                sonuc = yanit.text
                bolumadi = sayfalar[no]['ad']
                self.lblDersAd.setText(f"{bolumadi} {no+1}/{len(sayfalar)}")
                self.txtLink.setText( sayfalar[no]['link'] )
                self.lblStatus.setText(f"Durum= HTTP<{durum}>")
                self.txtDers.clear()
                if self.Kaydet:
                    paragraf = self.Belge.add_heading(bolumadi, level=1)
                    self.Paragraf = paragraf.add_run()
                soup = BeautifulSoup(sonuc, features='html.parser')
                # div = soup.find('div', {'id': 'sound'})
                # div = soup.select('div#iceriksayfa,div.icerik_sayfasi') # iki attribute'dan birini aramak için
                div = soup.find('div', {'class': 'icerik_sayfasi'})
                if not div:
                    div = soup.find('div', {'id': 'iceriksayfa'})
                if not div:
                    div = soup.find('div', {'id': 'icerik'})
                if not div:
                    div = soup.find('table', {'class': 'content'})
                if div:
                    metin = str(div.text)
                    metin = metin.replace('\n', ' ')
                    metin = metin.replace('\0', ' ')
                    self.txtDers.setPlainText(metin)
                    if self.Kaydet:
                        paragraf = self.Belge.add_paragraph(metin)
                        self.Paragraf = paragraf.add_run()
                imgs = soup.find_all('img')
                if imgs:
                    self.txtDers.appendHtml(self.imgGetir(imgs))
                if not div:
                    div = soup.find('div', {'id': 'guizno'})
                    if div:
                        html, metin = self.degerlendirmeSorulariGetir(soup, div.text)
                        self.txtDers.appendHtml( html )
                        if self.Kaydet: self.Belge.add_paragraph( metin )
                    else:
                        self.ctx.responseYaz(self.ctx.anaKlasor + f"{ICERIKKLASOR}\\oys-{self.ctx.IcerikDers[5:8]}no-{no}.html", sonuc)
                        div = soup.find('table')
                        if div:
                            sonuc = div.text
                        self.txtDers.appendHtml(sonuc)
                        if self.Kaydet: self.Belge.add_paragraph( '<HTML OKUNAMADI (flash vs olabilir)>' )
                if debug: print(f"IcerikOku: ders={bolumadi} status={durum} link={sayfalar[no]['link']}")
                return self.txtDers.toPlainText()

        def imgGetir(self, imgs, baseurl=''):
            html = ''
            for img in imgs:
                kaynak = baseurl + img['src']
                dosyaadi = kaynak.split('/')[-1]
                kaynak = self.ctx.adres + '/' + kaynak  #Viewer/temp/AYU/176/72/LRN//pics/2/bolum2.PNG
                new_kaynak = self.resimIndir( kaynak, dosyaadi)
                if debug: print(f"imgGetir img={img,img['src']} new_kaynak={new_kaynak}")
                new_img= img.text.replace(kaynak,new_kaynak)
                html += new_img
                if self.Kaydet:
                    en = int(img.get('width', '200')) - 20
                    boy = int(img.get('height', '150')) #- 10
                    if debug: print(f"imgGetir en={en} boy={boy}")
                    try:
                        self.Paragraf.add_picture(new_kaynak , width= Pt(en), height= Pt(boy) )
                    except:
                        self.Paragraf.add_text('<IMG hatalı>')
            return html

        def resimIndir(self, kaynak, dosyaadi):
            klasor = f"{self.ctx.anaKlasor}{ICERIKKLASOR}\\img"
            os.makedirs(klasor, exist_ok=True)
            new_kaynak = f"{klasor}\\{dosyaadi}"
            #print("new_kaynak=",new_kaynak)
            yanit = self.ctx.getSession().get( kaynak , cookies=self.ctx.cerezler, stream=True)
            with open(new_kaynak, 'wb') as dosya:
                for chunk in yanit.iter_content(chunk_size=512):
                    dosya.write(chunk)
            return new_kaynak

        def degerlendirmeSorulariGetir(self, soup, quizname):
            sorular=[]
            scripts = soup.find_all('script')
            kaynak = ''
            for script in scripts:
                script = script.text
                if 'var deger' in script:
                    kaynak = re.findall('deger=\"(.*?)\"', script)
                    if kaynak==[]:
                        kaynak = ''
                    else:
                        kaynak = kaynak[0]
            baseurl = '/' + kaynak[:-9]
            kaynak = self.ctx.adres + '/' + kaynak + quizname #Viewer/temp/AYU/700/912/LRN//sorular/quiz1.txt
            if debug: print (f"degerlendirmeSorulariGetir: baseurl={baseurl} kaynak={kaynak}")
            yanit = self.ctx.getSession().get(kaynak, cookies=self.ctx.cerezler)
            # if debug: print(f"degerlendirmeSorulariGetir:", yanit.encoding)
            yanit.encoding = 'utf-8'
            if yanit.status_code!=200: return yanit.status_code
            try:
                sonuc = yanit.json()
            except: #json.decoder.JSONDecodeError
                if debug: print(f"degerlendirmeSorulariGetir: yanit={yanit} kaynak={kaynak} status={yanit.status_code} text=\n{yanit.text}")
                return yanit.text
            if debug: print("degerlendirmeSorulariGetir: sonuc=", sonuc)
            quiz = sonuc['quiz']
            i = 0
            for soru in quiz:
                if soru['question'].find("img")>0:
                    imgs = []
                    soruimg = BeautifulSoup(soru['question'],features='html.parser').find('img')
                    if debug: print(f"degerlendirmeSorulariGetir: soruimg={soruimg}")
                    if soruimg:
                        imgs.append(soruimg)
                        self.imgGetir(imgs, baseurl )
                sorular.append({'soru': soru['question']})
                sorular[i]['siklar'] = soru['option']
                sorular[i]['cevap'] = soru['ans']
                i+=1
            html = ''
            metin = ''
            i = 0
            for soru in sorular:
                html += f"<b>Soru {i+1} : </b>" + soru['soru'] + '<br><ul>'
                metin += f"Soru {i+1} : " + soru['soru'] + '\n'
                # if debug: print(f"degerlendirmeSorulariGetir: siklar=", soru['siklar'])
                for opt, deger in soru['siklar'].items():
                    # if debug: print (f"degerlendirmeSorulariGetir: opt=", opt,"deger=",deger)
                    if deger.find("img") > 0:
                        imgs = []
                        soruimg = BeautifulSoup(deger, features='html.parser').find('img')
                        if debug: print(f"degerlendirmeSorulariGetir: şık.soruimg={soruimg}")
                        if soruimg:
                            imgs.append(soruimg)
                            self.imgGetir(imgs, baseurl )
                    if opt==soru['cevap']:
                        html += '<li><b>' + opt + ')' + deger + '</b></li>'
                        metin += '\n(' + opt + ') ' + deger
                    else:
                        html += '<li>' + opt + ')' + deger + '</li>'
                        metin +=  '\n ' + opt + ') ' + deger
                html += '</ul><br><br>'
                metin += '\n\n'
                i += 1
                # html += soru['cevap']
            return html, metin

if __name__ == '__main__':
    print('main.py çalıştır')

#INFO event işleme: QCoreApplication.processEvents()

